import streamlit as st
from langchain_together import ChatTogether
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import PromptTemplate
from langchain_core.documents import Document
import PyPDF2
import io
from docx import Document as DocxDocument


api_key = st.secrets["together_ai"]["api_key"]
model = st.secrets["together_ai"]["model"]

chat_model = ChatTogether(
    together_api_key=api_key,
    model=model
)

# Prompt engineering
prompt_template = PromptTemplate(
    input_variables=["context", "question"],
    template=(
        "You are an expert assistant. Use the provided context to answer questions accurately and concisely.\n"
        "Context: {context}\n\n"
        "Question: {question}\n\n"
        "Answer (be specific and avoid hallucinations):"
    )
)

# Define the QA chain
qa_chain = LLMChain(llm=chat_model, prompt=prompt_template)

st.title("📝 File Q&A")

# Initialize session state
if "messages" not in st.session_state:
    st.session_state["messages"] = [{"role": "assistant", "content": "Upload documents and ask questions about them!"}]
if "processed_files" not in st.session_state:
    st.session_state["processed_files"] = set()
if "file_contents" not in st.session_state:
    st.session_state["file_contents"] = {}


default_file_path = "FD.docx"  
try:
    with open(default_file_path, "rb") as f:
        file_content = None
        file_name = default_file_path.split("/")[-1]
        
        if default_file_path.endswith(".txt"):
            file_content = f.read().decode("utf-8")
        elif default_file_path.endswith(".pdf"):
            pdf_reader = PyPDF2.PdfReader(f)
            file_content = "".join(page.extract_text() for page in pdf_reader.pages)
        elif default_file_path.endswith(".docx"):
            doc = DocxDocument(f)
            file_content = "\n".join(para.text for para in doc.paragraphs)

        if file_content:
            document = Document(page_content=file_content, metadata={"source": file_name})
            chunk_size = 1000
            chunk_overlap = 200

            text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            chunks = text_splitter.split_documents([document])

            st.session_state["file_contents"][file_name] = "\n\n".join(chunk.page_content for chunk in chunks)
            st.success(f"Automatically loaded {file_name}")
except Exception as e:
    st.error(f"Error loading default file: {str(e)}")

uploaded_files = st.file_uploader("Upload articles", type=("txt", "pdf", "docx"), accept_multiple_files=True)

for uploaded_file in uploaded_files:
    if uploaded_file.name not in st.session_state["file_contents"]:
        try:
            if uploaded_file.type == "text/plain":
                file_content = uploaded_file.getvalue().decode("utf-8")
            elif uploaded_file.type == "application/pdf":
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.getvalue()))
                file_content = "".join(page.extract_text() for page in pdf_reader.pages)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = DocxDocument(io.BytesIO(uploaded_file.getvalue()))
                file_content = "\n".join(para.text for para in doc.paragraphs)
            else:
                st.error(f"Unsupported file type for {uploaded_file.name}. Please upload .txt, .pdf, or .docx files.")
                continue

            document = Document(page_content=file_content, metadata={"source": uploaded_file.name})
            chunk_size = 1000
            chunk_overlap = 200

            text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            chunks = text_splitter.split_documents([document])

            st.session_state["file_contents"][uploaded_file.name] = "\n\n".join(chunk.page_content for chunk in chunks)
            st.success(f"Processed {uploaded_file.name}")
        except Exception as e:
            st.error(f"Error processing {uploaded_file.name}: {str(e)}")

# Toggleable chat history
with st.expander("Chat History"):
    for msg in st.session_state.messages:
        if msg["role"] == "assistant":
            st.markdown(
                f"<div style='text-align: left; background-color: #f4f4f9; padding: 10px; border-radius: 10px; width: 80%; margin: 5px;'>Assistant: {msg['content']}</div>",
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                f"<div style='text-align: right; background-color: #d1e7fd; padding: 10px; border-radius: 10px; width: 80%; margin: 5px;'>User: {msg['content']}</div>",
                unsafe_allow_html=True,
            )

# Chat functionality
question = st.text_input("Ask a question about the uploaded documents:", disabled=len(st.session_state["file_contents"]) == 0)

if st.button("Send"):
    if question:
        st.session_state.messages.append({"role": "user", "content": question})
        st.chat_message("user").write(question)

        with st.chat_message("assistant"):
            try:
                all_contents = "\n\n".join(
                    [f"Content of {filename}:\n{content}" for filename, content in st.session_state["file_contents"].items()]
                )

                answer = qa_chain.run(context=all_contents, question=question)
                st.write(answer)

                st.session_state.messages.append({"role": "assistant", "content": answer})
            except Exception as e:
                st.error(f"Error querying Together AI model: {str(e)}")
